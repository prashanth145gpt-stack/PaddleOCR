import io
import re
from collections import Counter
from typing import Dict, List, Any

import cv2
import fitz
import numpy as np
from PIL import Image
import csv
from openpyxl import load_workbook
from docx import Document

from app.config import settings

def cfg(name, default):
    return getattr(settings, name, default)


SUCCESS = cfg("success_status", "SUCCESS")
FAILURE = cfg("failure_status", "FAILURE")


def is_pdf(file_bytes: bytes) -> bool:
    return file_bytes[:4] == b"%PDF"


def bytes_to_rgb_image(file_bytes: bytes):
    image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    return np.array(image)


def png_bytes_to_rgb_image(png_bytes: bytes):
    image = Image.open(io.BytesIO(png_bytes)).convert("RGB")
    return np.array(image)


def normalize_text(text: str) -> str:
    text = text or ""
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def get_text_stats(page) -> Dict[str, Any]:
    text = normalize_text(page.get_text("text"))
    blocks = page.get_text("blocks")
    text_blocks = []

    for block in blocks:
        if len(block) >= 5 and isinstance(block[4], str):
            block_text = normalize_text(block[4])
            if block_text:
                text_blocks.append(block_text)

    words = text.split()

    return {
        "char_count": int(len(text)),
        "word_count": int(len(words)),
        "text_block_count": int(len(text_blocks)),
        "sample_text": text[:200],
    }


def get_image_stats(doc, page) -> Dict[str, Any]:
    images = page.get_images(full=True)

    image_count = len(images)
    large_image_count = 0
    dominant_image_count = 0
    max_image_pixels = 0
    max_image_coverage_ratio = 0.0

    page_area = float(page.rect.width * page.rect.height)

    for img in images:
        xref = img[0]

        try:
            pix = fitz.Pixmap(doc, xref)
            pixels = int(pix.width * pix.height)
            max_image_pixels = max(max_image_pixels, pixels)

            if pixels > 500_000:
                large_image_count += 1

            rects = page.get_image_rects(xref)

            for rect in rects:
                image_area = float(rect.width * rect.height)
                coverage_ratio = image_area / page_area if page_area else 0.0
                max_image_coverage_ratio = max(max_image_coverage_ratio, coverage_ratio)

                if coverage_ratio >= 0.65:
                    dominant_image_count += 1

            pix = None

        except Exception:
            continue

    return {
        "image_count": int(image_count),
        "large_image_count": int(large_image_count),
        "dominant_image_count": int(dominant_image_count),
        "max_image_pixels": int(max_image_pixels),
        "max_image_coverage_ratio": round(float(max_image_coverage_ratio), 4),
    }


def classify_page(doc, page, page_number: int) -> Dict[str, Any]:
    text_stats = get_text_stats(page)
    image_stats = get_image_stats(doc, page)

    char_count = text_stats["char_count"]
    word_count = text_stats["word_count"]
    text_block_count = text_stats["text_block_count"]

    image_count = image_stats["image_count"]
    dominant_image_count = image_stats["dominant_image_count"]

    has_text_layer = (
        char_count >= settings.text_min_chars_per_page
        and word_count >= settings.text_min_words_per_page
    )

    has_substantial_text = (
        char_count >= settings.text_good_chars_per_page
        and word_count >= settings.text_good_words_per_page
        and text_block_count > 0
    )

    has_images = image_count > 0
    has_dominant_image = dominant_image_count > 0

    if has_substantial_text and not has_dominant_image:
        page_type = "DIGITAL_TEXT_PAGE"
    elif has_text_layer and has_dominant_image:
        page_type = "OCR_LAYERED_SCAN_PAGE"
    elif not has_text_layer and has_images:
        page_type = "SCANNED_IMAGE_PAGE"
    elif char_count < settings.text_min_chars_per_page and image_count == 0:
        page_type = "BLANK_OR_LOW_CONTENT_PAGE"
    else:
        page_type = "UNKNOWN_PAGE"

    return {
        "page_number": int(page_number),
        "page_type": page_type,
        "text_stats": text_stats,
        "image_stats": image_stats,
    }


def classify_pdf(file_bytes: bytes) -> Dict[str, Any]:
    doc = fitz.open(stream=file_bytes, filetype="pdf")

    try:
        if doc.needs_pass:
            raise ValueError("Password-protected PDFs are not supported.")

        total_pages = len(doc)

        if total_pages == 0:
            raise ValueError("PDF has no pages.")

        if total_pages > settings.max_allowed_pages:
            raise ValueError(
                f"PDF has too many pages. Maximum allowed pages: {settings.max_allowed_pages}."
            )

        pages = []

        for index in range(total_pages):
            pages.append(classify_page(doc, doc[index], index + 1))

        counts = Counter(page["page_type"] for page in pages)

        digital_pages = counts["DIGITAL_TEXT_PAGE"]
        scanned_pages = counts["SCANNED_IMAGE_PAGE"]
        ocr_layered_pages = counts["OCR_LAYERED_SCAN_PAGE"]

        digital_ratio = digital_pages / total_pages
        scanned_ratio = scanned_pages / total_pages
        ocr_layered_ratio = ocr_layered_pages / total_pages

        if digital_ratio >= settings.machine_readable_ratio and scanned_pages == 0 and ocr_layered_pages == 0:
            document_type = "MACHINE_READABLE_PDF"
        elif scanned_ratio >= settings.scanned_ratio and digital_pages == 0:
            document_type = "SCANNED_PDF"
        elif ocr_layered_ratio >= settings.ocr_layered_ratio:
            document_type = "SEARCHABLE_SCANNED_PDF"
        elif digital_pages > 0 and (scanned_pages > 0 or ocr_layered_pages > 0):
            document_type = "MIXED_PDF"
        else:
            document_type = "UNKNOWN_OR_LOW_CONTENT_PDF"

        return {
            "document_type": document_type,
            "total_pages": int(total_pages),
            "page_type_counts": dict(counts),
            "pages": pages,
        }

    finally:
        doc.close()

def render_pdf_page(file_bytes: bytes, page_index: int, dpi: int = None) -> bytes:
    dpi = dpi or settings.render_dpi
    doc = fitz.open(stream=file_bytes, filetype="pdf")

    try:
        if page_index < 0 or page_index >= len(doc):
            raise ValueError("Invalid PDF page index.")

        page = doc[page_index]
        pix = page.get_pixmap(dpi=dpi, alpha=False)
        return pix.tobytes("png")

    finally:
        doc.close()


def get_validation_strategy(page_type: str) -> str:
    if page_type == "DIGITAL_TEXT_PAGE":
        return "TEXT_LAYER_VALIDATION"
    if page_type == "SCANNED_IMAGE_PAGE":
        return "OCR_IMAGE_QUALITY_VALIDATION"
    if page_type == "OCR_LAYERED_SCAN_PAGE":
        return "HYBRID_TEXT_AND_IMAGE_VALIDATION"
    if page_type == "BLANK_OR_LOW_CONTENT_PAGE":
        return "BLANK_PAGE_VALIDATION"
    if page_type == "IMAGE_DOCUMENT":
        return "OCR_IMAGE_QUALITY_VALIDATION"
    return "UNKNOWN_PAGE_VALIDATION"


def normalize_score(value: float, min_value: float, max_value: float) -> float:
    if value <= min_value:
        return 0.0
    if value >= max_value:
        return 100.0
    return ((value - min_value) / (max_value - min_value)) * 100.0


def calculate_blur(gray_image) -> Dict[str, Any]:
    lap_var = cv2.Laplacian(gray_image, cv2.CV_64F).var()
    return {
        "laplacian_variance": float(lap_var),
        "blur_score": round(normalize_score(lap_var, 20, 250), 2),
    }


def calculate_contrast(gray_image) -> Dict[str, Any]:
    contrast_std = float(gray_image.std())
    return {
        "contrast_std": contrast_std,
        "contrast_score": round(normalize_score(contrast_std, 20, 80), 2),
    }


def calculate_brightness(gray_image) -> Dict[str, Any]:
    brightness_mean = float(gray_image.mean())

    if settings.min_brightness <= brightness_mean <= settings.max_brightness:
        score = 100.0
    else:
        distance = (
            settings.min_brightness - brightness_mean
            if brightness_mean < settings.min_brightness
            else brightness_mean - settings.max_brightness
        )
        score = max(0.0, 100.0 - distance)

    return {
        "brightness_mean": brightness_mean,
        "brightness_score": round(score, 2),
    }


def calculate_resolution(image_rgb) -> Dict[str, Any]:
    height, width = image_rgb.shape[:2]

    width_score = normalize_score(width, 600, settings.min_width)
    height_score = normalize_score(height, 600, settings.min_height)

    return {
        "width": int(width),
        "height": int(height),
        "resolution_score": round(min(width_score, height_score), 2),
    }


def calculate_noise(gray_image) -> Dict[str, Any]:
    median_image = cv2.medianBlur(gray_image, 3)
    noise_value = np.mean(
        np.abs(gray_image.astype("float") - median_image.astype("float"))
    )

    return {
        "noise_value": round(float(noise_value), 2),
        "noise_score": round(max(0.0, 100.0 - noise_value * 3.0), 2),
    }


def estimate_skew(gray_image) -> Dict[str, Any]:
    edges = cv2.Canny(gray_image, 50, 150, apertureSize=3)
    lines = cv2.HoughLines(edges, 1, np.pi / 180, 150)

    if lines is None:
        return {"skew_angle": 0.0, "skew_score": 100.0}

    angles = []

    for line in lines[:100]:
        _, theta = line[0]
        angle = (theta * 180 / np.pi) - 90

        if -45 < angle < 45:
            angles.append(angle)

    if not angles:
        return {"skew_angle": 0.0, "skew_score": 100.0}

    skew_angle = float(np.median(angles))

    return {
        "skew_angle": round(skew_angle, 2),
        "skew_score": round(max(0.0, 100.0 - abs(skew_angle) * 10.0), 2),
    }


def detect_blank_page(gray_image) -> bool:
    return bool(gray_image.std() < 10 and gray_image.mean() > 220)


def detect_xerox_like_scan(gray_image, contrast_std: float, noise_value: float) -> bool:
    edges = cv2.Canny(gray_image, 80, 200)
    edge_density = float(np.count_nonzero(edges) / edges.size)

    return bool(
        float(contrast_std) < 32
        and float(noise_value) > 12
        and edge_density < 0.08
    )


def image_issue(code, message, severity):
    return {
        "source": "IMAGE_QUALITY",
        "code": code,
        "message": message,
        "severity": severity,
    }


def analyze_image_quality(image_rgb) -> Dict[str, Any]:
    gray_image = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2GRAY)

    blur = calculate_blur(gray_image)
    contrast = calculate_contrast(gray_image)
    brightness = calculate_brightness(gray_image)
    resolution = calculate_resolution(image_rgb)
    noise = calculate_noise(gray_image)
    skew = estimate_skew(gray_image)

    is_blank = detect_blank_page(gray_image)
    is_xerox_like = detect_xerox_like_scan(
        gray_image,
        contrast["contrast_std"],
        noise["noise_value"],
    )

    issues = []

    if is_blank:
        issues.append(image_issue("BLANK_PAGE", "The page appears blank or nearly blank.", "CRITICAL"))

    laplacian_variance = blur["laplacian_variance"]

    if laplacian_variance < settings.unreadable_blur_threshold:
        issues.append(
            image_issue(
                "UNREADABLE_BLUR",
                "The page is too blurred for reliable OCR readability.",
                "CRITICAL"
            )
        )

    elif laplacian_variance < settings.severe_blur_threshold:
        issues.append(
            image_issue(
                "SEVERE_BLUR_RECOVERABLE",
                "The page is severely blurred, but OCR readability will be checked.",
                "OCR_RECOVERY_REQUIRED"
            )
        )

    elif laplacian_variance < settings.min_laplacian_variance:
        issues.append(
            image_issue(
                "BLUR_DETECTED",
                "The page appears blurred.",
                "WARNING"
            )
        )

    if contrast["contrast_std"] < settings.min_contrast_std:
        issues.append(image_issue("POOR_CONTRAST", "The page has poor contrast between text and background.", "WARNING"))

    if resolution["resolution_score"] < 30:
        issues.append(image_issue("LOW_RESOLUTION", "The page resolution is extremely low.", "CRITICAL"))
    elif resolution["resolution_score"] < 60:
        issues.append(image_issue("LOW_RESOLUTION", "The page resolution is low.", "WARNING"))

    if not settings.min_brightness <= brightness["brightness_mean"] <= settings.max_brightness:
        issues.append(image_issue("BAD_EXPOSURE", "The page is too bright or too dark.", "WARNING"))

    if abs(skew["skew_angle"]) > settings.max_skew_angle:
        issues.append(image_issue("SKEW_DETECTED", "The page appears tilted.", "WARNING"))

    if is_xerox_like:
        issues.append(image_issue("XEROX_COPY_LIKE_SCAN", "The page appears like a poor-quality photocopy scan.", "WARNING"))

    quality_score = (
        blur["blur_score"] * 0.20
        + resolution["resolution_score"] * 0.15
        + contrast["contrast_score"] * 0.15
        + brightness["brightness_score"] * 0.10
        + noise["noise_score"] * 0.10
        + skew["skew_score"] * 0.10
        + 20.0
    )

    metrics = {
        "blur_laplacian_variance": round(float(blur["laplacian_variance"]), 2),
        "blur_score": float(blur["blur_score"]),
        "contrast_std": round(float(contrast["contrast_std"]), 2),
        "contrast_score": float(contrast["contrast_score"]),
        "brightness_mean": round(float(brightness["brightness_mean"]), 2),
        "brightness_score": float(brightness["brightness_score"]),
        "width": int(resolution["width"]),
        "height": int(resolution["height"]),
        "resolution_score": float(resolution["resolution_score"]),
        "noise_value": float(noise["noise_value"]),
        "noise_score": float(noise["noise_score"]),
        "skew_angle": float(skew["skew_angle"]),
        "skew_score": float(skew["skew_score"]),
        "is_blank": bool(is_blank),
        "is_xerox_like": bool(is_xerox_like),
    }

    quality_score = min(100.0, max(0.0, quality_score))

    critical_codes = {
        issue.get("code")
        for issue in issues
        if issue.get("severity") == "CRITICAL"
    }

    if "BLANK_PAGE" in critical_codes:
        quality_score = 0.0

    elif "SEVERE_BLUR" in critical_codes:
        quality_score = min(quality_score, 40.0)

    elif "LOW_RESOLUTION" in critical_codes:
        quality_score = min(quality_score, 55.0)

    return {
        "quality_score": round(quality_score, 2),
        "issues": issues,
        "metrics": metrics,
    }


def has_critical_issue(issues: List[Dict[str, Any]]) -> bool:
    return any(issue.get("severity") == "CRITICAL" for issue in issues)

def has_blocking_critical_issue(issues: List[Dict[str, Any]]) -> bool:
    """
    Returns True only for issues that should directly block the page
    without OCR recovery.
    """
    return any(
        issue.get("severity") == "CRITICAL"
        for issue in issues
    )


def requires_ocr_recovery(issues: List[Dict[str, Any]]) -> bool:
    """
    Returns True when image quality is poor but OCR should still be attempted
    to prove readability.
    """
    if not settings.allow_ocr_recovery_for_severe_blur:
        return False

    return any(
        issue.get("code") == "SEVERE_BLUR_RECOVERABLE"
        for issue in issues
    )


def ocr_recovery_passed(ocr_metrics: Dict[str, Any]) -> bool:
    """
    Stronger OCR pass condition used only when we are trying to recover
    a page that has recoverable severe blur.
    """
    if ocr_metrics.get("ocr_error"):
        return False

    return (
        ocr_metrics.get("mean_confidence", 0.0) >= settings.ocr_recovery_mean_confidence
        and ocr_metrics.get("median_confidence", 0.0) >= settings.ocr_recovery_median_confidence
        and ocr_metrics.get("detected_text_boxes", 0) >= settings.ocr_recovery_min_text_boxes
        and ocr_metrics.get("detected_chars", 0) >= settings.ocr_recovery_min_detected_chars
        and ocr_metrics.get("low_confidence_line_ratio", 1.0) <= settings.ocr_recovery_max_low_conf_ratio
    )

def has_strong_text_layer(page_info: Dict[str, Any]) -> bool:
    text_stats = page_info["text_stats"]

    return (
        text_stats["char_count"] >= cfg("min_searchable_text_chars_per_page", 100)
        and text_stats["word_count"] >= cfg("min_searchable_text_words_per_page", 15)
        and text_stats["text_block_count"] > 0
    )


def text_issue(code, message, severity="CRITICAL"):
    return {
        "source": "TEXT_LAYER",
        "code": code,
        "message": message,
        "severity": severity,
    }


def ocr_issue(code, message, severity="CRITICAL"):
    return {
        "source": "OCR_READABILITY",
        "code": code,
        "message": message,
        "severity": severity,
    }


def decide_text_layer_page(page_info: Dict[str, Any]) -> Dict[str, Any]:
    text_stats = page_info["text_stats"]
    issues = []

    if text_stats["char_count"] < settings.text_min_chars_per_page:
        issues.append(text_issue("INSUFFICIENT_TEXT_LAYER", "The page does not contain enough extractable text."))

    if text_stats["word_count"] < settings.text_min_words_per_page:
        issues.append(text_issue("LOW_WORD_COUNT", "The page has very low readable text content."))

    if text_stats["text_block_count"] == 0:
        issues.append(text_issue("NO_TEXT_BLOCKS", "No meaningful text blocks were detected."))

    if issues:
        return {
            "status": FAILURE,
            "quality_score": 40.0,
            "issues": issues,
            "warnings": [],
            "metrics": {
                "text_stats": text_stats,
                "image_stats": page_info["image_stats"],
            },
        }

    return {
        "status": SUCCESS,
        "quality_score": 95.0,
        "issues": [],
        "warnings": [],
        "metrics": {
            "text_stats": text_stats,
            "image_stats": page_info["image_stats"],
        },
    }


def decide_blank_or_low_content_page(page_info: Dict[str, Any], total_pages: int) -> Dict[str, Any]:
    if total_pages <= 1:
        return {
            "status": FAILURE,
            "quality_score": 0.0,
            "issues": [
                text_issue(
                    "BLANK_OR_LOW_CONTENT_DOCUMENT",
                    "The uploaded document appears blank or has insufficient readable content.",
                )
            ],
            "warnings": [],
            "metrics": {
                "text_stats": page_info["text_stats"],
                "image_stats": page_info["image_stats"],
            },
        }

    return {
        "status": SUCCESS,
        "quality_score": 90.0,
        "issues": [],
        "warnings": [
            {
                "source": "TEXT_LAYER",
                "code": "LOW_CONTENT_PAGE",
                "message": "Page appears blank, decorative, or image-only, but the overall document is multi-page.",
                "severity": "WARNING",
            }
        ],
        "metrics": {
            "text_stats": page_info["text_stats"],
            "image_stats": page_info["image_stats"],
        },
    }


def decide_searchable_scanned_page(page_info: Dict[str, Any], image_quality_result: Dict[str, Any]) -> Dict[str, Any]:
    text_strong = has_strong_text_layer(page_info)
    image_issues = image_quality_result["issues"]
    critical_issues = [issue for issue in image_issues if issue.get("severity") == "CRITICAL"]

    warnings = [
    issue for issue in image_issues
    if issue.get("severity") != "CRITICAL"]

    if not text_strong:
        return {
            "status": FAILURE,
            "quality_score": image_quality_result["quality_score"],
            "issues": [
                text_issue(
                    "WEAK_EMBEDDED_TEXT_LAYER",
                    "Embedded text layer is weak or insufficient for reliable readability.",
                )
            ],
            "warnings": warnings,
            "metrics": {
                "text_stats": page_info["text_stats"],
                "image_stats": page_info["image_stats"],
                "image_quality": image_quality_result["metrics"],
                "ocr": {
                    "skipped": True,
                    "reason": "OCR not run because page is searchable scanned PDF.",
                },
            },
        }

    if critical_issues:
        return {
            "status": FAILURE,
            "quality_score": image_quality_result["quality_score"],
            "issues": critical_issues,
            "warnings": warnings,
            "metrics": {
                "text_stats": page_info["text_stats"],
                "image_stats": page_info["image_stats"],
                "image_quality": image_quality_result["metrics"],
                "ocr": {
                    "skipped": True,
                    "reason": "Embedded text layer exists but image has critical quality issue.",
                },
            },
        }

    return {
        "status": SUCCESS,
        "quality_score": max(85.0, image_quality_result["quality_score"]),
        "issues": [],
        "warnings": warnings,
        "metrics": {
            "text_stats": page_info["text_stats"],
            "image_stats": page_info["image_stats"],
            "image_quality": image_quality_result["metrics"],
            "ocr": {
                "skipped": True,
                "reason": "Embedded text layer is strong. OCR was not required.",
            },
        },
    }


def decide_image_quality_only_page(page_info: Dict[str, Any], image_quality_result: Dict[str, Any]) -> Dict[str, Any]:
    issues = image_quality_result["issues"]
    critical_issues = [issue for issue in issues if issue.get("severity") == "CRITICAL"]
    warnings = [issue for issue in issues if issue.get("severity") != "CRITICAL"]

    if critical_issues:
        return {
            "status": FAILURE,
            "quality_score": image_quality_result["quality_score"],
            "issues": critical_issues,
            "warnings": warnings,
            "metrics": {
                "text_stats": page_info["text_stats"],
                "image_stats": page_info["image_stats"],
                "image_quality": image_quality_result["metrics"],
                "ocr": {
                    "skipped": True,
                    "reason": "OCR skipped because image quality has critical issues.",
                },
            },
        }

    if image_quality_result["quality_score"] < 60:
        warnings.append(
            image_issue(
                "LOW_IMAGE_QUALITY_SCORE",
                "The page image quality is borderline.",
                "WARNING",
            )
        )

    return {
        "status": SUCCESS,
        "quality_score": image_quality_result["quality_score"],
        "issues": [],
        "warnings": warnings,
        "metrics": {
            "text_stats": page_info["text_stats"],
            "image_stats": page_info["image_stats"],
            "image_quality": image_quality_result["metrics"],
            "ocr": {
                "skipped": True,
                "reason": "Page was not selected for OCR sampling.",
            },
        },
    }


def decide_ocr_page(
    image_quality_result: Dict[str, Any],
    ocr_metrics: Dict[str, Any],
    page_type: str,
    total_pages: int) -> Dict[str, Any]:

    image_issues = image_quality_result["issues"]

    critical_image_issues = [
        issue for issue in image_issues
        if issue.get("severity") == "CRITICAL"
    ]

    warnings = [
        issue for issue in image_issues
        if issue.get("severity") != "CRITICAL"
    ]

    issues = list(critical_image_issues)

    mean_conf = ocr_metrics.get("mean_confidence", 0.0)
    median_conf = ocr_metrics.get("median_confidence", 0.0)
    low_ratio = ocr_metrics.get("low_confidence_line_ratio", 1.0)
    boxes = ocr_metrics.get("detected_text_boxes", 0)
    chars = ocr_metrics.get("detected_chars", 0)

    ocr_failures = []

    if ocr_metrics.get("ocr_error"):
        ocr_failures.append(
            ocr_issue(
                "OCR_ENGINE_ERROR",
                f"OCR engine failed: {ocr_metrics.get('ocr_error')}"
            )
    )

    if boxes < settings.min_text_boxes:
        ocr_failures.append(
            ocr_issue(
                "TOO_FEW_TEXT_BOXES",
                "Too few readable text regions were detected."
            )
        )

    if chars < settings.min_detected_chars:
        ocr_failures.append(
            ocr_issue(
                "TOO_FEW_READABLE_CHARACTERS",
                "Too few readable characters were detected."
            )
        )

    if mean_conf < settings.ocr_reject_mean_confidence:
        ocr_failures.append(
            ocr_issue(
                "LOW_MEAN_OCR_CONFIDENCE",
                "OCR confidence is too low for reliable reading."
            )
        )

    if median_conf < settings.ocr_reject_median_confidence:
        ocr_failures.append(
            ocr_issue(
                "LOW_MEDIAN_OCR_CONFIDENCE",
                "Median OCR confidence is too low for reliable reading."
            )
        )

    if low_ratio > settings.max_low_conf_ratio_reject:
        ocr_failures.append(
            ocr_issue(
                "TOO_MANY_LOW_CONFIDENCE_LINES",
                "Too many detected text lines have low OCR confidence."
            )
        )

    image_score = image_quality_result["quality_score"]
    ocr_score = mean_conf * 100.0
    final_score = round(image_score * 0.60 + ocr_score * 0.40, 2)

    if critical_image_issues:
        return {
            "status": FAILURE,
            "quality_score": final_score,
            "issues": issues,
            "warnings": warnings,
            "metrics": {
                "image_quality": image_quality_result["metrics"],
                "ocr": ocr_metrics,
            },
        }

    if (
        page_type == "UNKNOWN_PAGE"
        and total_pages > 1
        and settings.unknown_page_ocr_failure_is_warning
        and ocr_failures
    ):
        for item in ocr_failures:
            warning = dict(item)
            warning["severity"] = "WARNING"
            warning["code"] = "LOW_TEXT_DENSITY_ON_UNKNOWN_PAGE"
            warning["message"] = (
                "This unknown page has low OCR text density, but it is not treated "
                "as a blocking failure because the document has multiple pages."
            )
            warnings.append(warning)

        return {
            "status": SUCCESS,
            "quality_score": max(70.0, final_score),
            "issues": [],
            "warnings": warnings,
            "metrics": {
                "image_quality": image_quality_result["metrics"],
                "ocr": ocr_metrics,
            },
        }

    if ocr_failures:
        issues.extend(ocr_failures)

        return {
            "status": FAILURE,
            "quality_score": final_score,
            "issues": issues,
            "warnings": warnings,
            "metrics": {
                "image_quality": image_quality_result["metrics"],
                "ocr": ocr_metrics,
            },
        }

    return {
        "status": SUCCESS,
        "quality_score": final_score,
        "issues": [],
        "warnings": warnings,
        "metrics": {
            "image_quality": image_quality_result["metrics"],
            "ocr": ocr_metrics,
        },
    }

def decide_ocr_recovery_page(
    image_quality_result: Dict[str, Any],
    ocr_metrics: Dict[str, Any],
) -> Dict[str, Any]:
    """
    Handles pages where image blur is severe but potentially recoverable.
    If OCR proves strong readability, page passes with warning.
    If OCR fails, page fails.
    """
    recovery_passed = ocr_recovery_passed(ocr_metrics)

    if recovery_passed:
        return {
            "status": SUCCESS,
            "quality_score": max(
                settings.min_document_quality_score,
                min(image_quality_result["quality_score"], 82.0)
            ),
            "issues": [],
            "warnings": [
                {
                    "source": "IMAGE_QUALITY",
                    "code": "SEVERE_BLUR_BUT_OCR_READABLE",
                    "message": "The page appears severely blurred, but OCR readability check passed.",
                    "severity": "WARNING",
                }
            ],
            "metrics": {
                "image_quality": image_quality_result["metrics"],
                "ocr": ocr_metrics,
            },
        }

    return {
        "status": FAILURE,
        "quality_score": min(image_quality_result["quality_score"], 40.0),
        "issues": [
            {
                "source": "IMAGE_QUALITY",
                "code": "SEVERE_BLUR",
                "message": "The page is severely blurred and OCR readability check failed.",
                "severity": "CRITICAL",
            },
            ocr_issue(
                "OCR_RECOVERY_FAILED",
                "OCR could not reliably read the severely blurred page."
            )
        ],
        "warnings": [],
        "metrics": {
            "image_quality": image_quality_result["metrics"],
            "ocr": ocr_metrics,
        },
    }

def decide_document(page_results: List[Dict[str, Any]], total_pages: int) -> Dict[str, Any]:
    if not page_results:
        return {
            "status": FAILURE,
            "document_quality_score": 0.0,
            "is_readable": False,
            "reupload_required": True,
            "message": "No readable pages were found.",
        }

    score = round(
        sum(page["quality_score"] for page in page_results) / len(page_results),
        2
    )

    failed_pages = [
        page for page in page_results
        if page["status"] == FAILURE
    ]

    critical_non_ocr_failures = []
    ocr_failure_pages = []

    for page in failed_pages:
        issues = page.get("issues", [])

        has_ocr_failure = any(
            issue.get("source") == "OCR_READABILITY"
            for issue in issues
        )

        has_non_ocr_critical = any(
            issue.get("severity") == "CRITICAL"
            and issue.get("source") != "OCR_READABILITY"
            for issue in issues
        )

        if has_non_ocr_critical:
            critical_non_ocr_failures.append(page)

        if has_ocr_failure:
            ocr_failure_pages.append(page)

    low_content_warnings = [
        warning
        for page in page_results
        for warning in page.get("warnings", [])
        if warning.get("code") == "LOW_CONTENT_PAGE"
    ]

    low_content_ratio = len(low_content_warnings) / total_pages if total_pages else 0

    if critical_non_ocr_failures:
        return {
            "status": FAILURE,
            "document_quality_score": score,
            "is_readable": False,
            "reupload_required": True,
            "message": "Document quality is not acceptable. Please re-upload a clearer document.",
        }

    if total_pages <= 1 and ocr_failure_pages:
        return {
            "status": FAILURE,
            "document_quality_score": score,
            "is_readable": False,
            "reupload_required": True,
            "message": "Document quality is not acceptable. Please re-upload a clearer document.",
        }

    ocr_failure_ratio = (
        len(ocr_failure_pages) / max(1, len(page_results))
    )

    if (
        len(ocr_failure_pages) > settings.max_ocr_failure_pages_allowed
        or ocr_failure_ratio > settings.max_ocr_failure_ratio_allowed
    ):
        return {
            "status": FAILURE,
            "document_quality_score": score,
            "is_readable": False,
            "reupload_required": True,
            "message": "Document quality is not acceptable. Please re-upload a clearer document.",
        }

    if (
        len(low_content_warnings) > settings.max_low_content_pages_allowed
        and low_content_ratio > settings.max_low_content_page_ratio
    ):
        return {
            "status": FAILURE,
            "document_quality_score": score,
            "is_readable": False,
            "reupload_required": True,
            "message": "Too many pages appear blank or have insufficient readable content.",
        }

    if score < settings.min_document_quality_score:
        return {
            "status": FAILURE,
            "document_quality_score": score,
            "is_readable": False,
            "reupload_required": True,
            "message": (
                f"Document quality score is below the minimum acceptable threshold "
                f"of {settings.min_document_quality_score}. Please re-upload a clearer document."
            ),
        }
    
    return {
        "status": SUCCESS,
        "document_quality_score": score,
        "is_readable": True,
        "reupload_required": False,
        "message": "Document quality is acceptable for submission.",
    }

def make_json_safe(value):
    if isinstance(value, dict):
        return {str(key): make_json_safe(val) for key, val in value.items()}
    if isinstance(value, list):
        return [make_json_safe(item) for item in value]
    if isinstance(value, tuple):
        return [make_json_safe(item) for item in value]
    if isinstance(value, np.bool_):
        return bool(value)
    if isinstance(value, np.integer):
        return int(value)
    if isinstance(value, np.floating):
        return float(value)
    if isinstance(value, np.ndarray):
        return value.tolist()
    return value


def select_ocr_candidate_pages(scanned_candidates: List[Dict[str, Any]], max_ocr_pages: int) -> List:
    if not scanned_candidates:
        return []

    selected = []
    selected_set = set()

    def add(candidate):
        page_number = candidate["page_number"]
        if page_number not in selected_set and len(selected) < max_ocr_pages:
            selected.append(page_number)
            selected_set.add(page_number)

    ocr_recovery_pages = [
        c for c in scanned_candidates
        if c.get("ocr_recovery_required")
    ]

    unknown_pages = [
        c for c in scanned_candidates
        if c["page_type"] == "UNKNOWN_PAGE"
    ]

    scanned_pages = [
        c for c in scanned_candidates
        if c["page_type"] == "SCANNED_IMAGE_PAGE"
    ]

    for c in ocr_recovery_pages:
        add(c)

    for c in unknown_pages:
        add(c)

    for c in scanned_pages[:3]:
        add(c)

    for c in scanned_pages[-2:]:
        add(c)

    worst_pages = sorted(
        scanned_candidates,
        key=lambda x: x["image_quality_result"]["quality_score"],
    )

    for c in worst_pages:
        add(c)

    return selected

def summarize_warnings(warnings: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    grouped = {}

    for warning in warnings:
        key = (
            warning.get("source", "UNKNOWN"),
            warning.get("code", "UNKNOWN"),
            warning.get("message", "")
        )

        if key not in grouped:
            grouped[key] = {
                "source": warning.get("source", "UNKNOWN"),
                "code": warning.get("code", "UNKNOWN"),
                "message": warning.get("message", ""),
                "severity": warning.get("severity", "WARNING"),
                "affected_page_count": 0,
                "sample_page_numbers": []
            }

        grouped[key]["affected_page_count"] += 1

        page_number = warning.get("page_number")

        if (
            page_number is not None
            and len(grouped[key]["sample_page_numbers"]) < settings.warning_sample_pages_limit
        ):
            grouped[key]["sample_page_numbers"].append(page_number)

    return list(grouped.values())

def summarize_failures(page_results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    grouped = {}

    for page in page_results:
        if page.get("status") != FAILURE:
            continue

        page_number = page.get("page_number")

        for issue in page.get("issues", []):
            key = (
                issue.get("source", "UNKNOWN"),
                issue.get("code", "UNKNOWN"),
                issue.get("message", "")
            )

            if key not in grouped:
                grouped[key] = {
                    "source": issue.get("source", "UNKNOWN"),
                    "code": issue.get("code", "UNKNOWN"),
                    "message": issue.get("message", ""),
                    "severity": issue.get("severity", "CRITICAL"),
                    "affected_page_count": 0,
                    "page_numbers": []
                }

            grouped[key]["affected_page_count"] += 1

            if page_number is not None:
                grouped[key]["page_numbers"].append(page_number)

    for item in grouped.values():
        item["page_numbers"] = sorted(list(set(item["page_numbers"])))

    return list(grouped.values())

def build_ui_response(
    document_decision: Dict[str, Any],
    document_type: str,
    validation_strategy: str,
    total_pages: int,
    page_results: List[Dict[str, Any]],
    page_type_counts: Dict[str, int],
) -> Dict[str, Any]:
    warnings = []

    for page in page_results:
        for warning in page.get("warnings", []):
            item = dict(warning)
            item["page_number"] = page["page_number"]
            warnings.append(item)

    response = {
        "status": document_decision["status"],
        "document_type": document_type,
        "validation_strategy": validation_strategy,
        "document_quality_score": document_decision["document_quality_score"],
        "is_readable": document_decision["is_readable"],
        "reupload_required": document_decision["reupload_required"],
        "message": document_decision["message"],
        "total_pages": total_pages,
        "metadata": {
            "page_type_counts": page_type_counts
        }
    }

    validation_failures = summarize_failures(page_results)

    if validation_failures:
        response["validation_failures"] = validation_failures

    if warnings and settings.return_warnings_summary:
        response["warnings_summary"] = summarize_warnings(warnings)

    if warnings and settings.return_warning_details:
        response["warnings"] = warnings

    if settings.return_success_pages:
        response["pages"] = page_results

    return make_json_safe(response)

def is_word_document(file_name: str = "") -> bool:
    file_name = (file_name or "").lower()
    return file_name.endswith((".docx", ".doc"))

def is_spreadsheet(file_name: str = "") -> bool:
    file_name = (file_name or "").lower()
    return file_name.endswith((".xlsx", ".xls", ".csv"))


def spreadsheet_failure(code: str, message: str, file_name: str = "") -> Dict[str, Any]:
    return {
        "status": FAILURE,
        "document_type": "SPREADSHEET",
        "validation_strategy": "SPREADSHEET_STRUCTURE_VALIDATION",
        "document_quality_score": 0.0,
        "is_readable": False,
        "reupload_required": True,
        "message": "Spreadsheet quality validation failed. Please re-upload a readable file.",
        "total_pages": 0,
        "validation_failures": [
            {
                "source": "SPREADSHEET_VALIDATION",
                "code": code,
                "message": message,
                "severity": "CRITICAL"
            }
        ],
        "metadata": {
            "file_name": file_name
        }
    }

def word_failure(code: str, message: str, file_name: str = "") -> Dict[str, Any]:
    return {
        "status": FAILURE,
        "document_type": "WORD_DOCUMENT",
        "validation_strategy": "WORD_DOCUMENT_STRUCTURE_VALIDATION",
        "document_quality_score": 0.0,
        "is_readable": False,
        "reupload_required": True,
        "message": "Word document quality validation failed. Please re-upload a readable file.",
        "total_pages": 0,
        "validation_failures": [
            {
                "source": "WORD_DOCUMENT_VALIDATION",
                "code": code,
                "message": message,
                "severity": "CRITICAL"
            }
        ],
        "metadata": {
            "file_name": file_name
        }
    }

def validate_spreadsheet(file_bytes: bytes, file_name: str = "") -> Dict[str, Any]:
    file_name_lower = (file_name or "").lower()

    try:
        if file_name_lower.endswith(".csv"):
            try:
                text = file_bytes.decode("utf-8-sig")
            except UnicodeDecodeError:
                return spreadsheet_failure(
                    "UNREADABLE_CSV_ENCODING",
                    "CSV file encoding is not readable as UTF-8.",
                    file_name
                )

            rows = list(csv.reader(io.StringIO(text)))

            non_empty_rows = [
                row for row in rows
                if any(str(cell).strip() for cell in row)
            ]

            if not non_empty_rows:
                return spreadsheet_failure(
                    "EMPTY_SPREADSHEET",
                    "CSV file does not contain readable data.",
                    file_name
                )

            return {
                "status": SUCCESS,
                "document_type": "SPREADSHEET",
                "validation_strategy": "SPREADSHEET_STRUCTURE_VALIDATION",
                "document_quality_score": 100.0,
                "is_readable": True,
                "reupload_required": False,
                "message": "Spreadsheet is readable and acceptable for submission.",
                "total_pages": 0,
                "metadata": {
                    "file_name": file_name,
                    "file_type": "CSV",
                    "non_empty_rows": len(non_empty_rows)
                }
            }

        if file_name_lower.endswith(".xls"):
            return spreadsheet_failure(
                "UNSUPPORTED_LEGACY_EXCEL",
                "Legacy .xls files are not supported in this validator. Please upload .xlsx or CSV.",
                file_name
            )

        workbook = load_workbook(
            io.BytesIO(file_bytes),
            read_only=True,
            data_only=True
        )

        sheet_count = len(workbook.sheetnames)
        non_empty_sheets = 0
        sampled_non_empty_cells = 0

        for sheet_name in workbook.sheetnames:
            ws = workbook[sheet_name]
            has_data = False

            for row in ws.iter_rows(max_row=100, max_col=50, values_only=True):
                for cell in row:
                    if cell is not None and str(cell).strip():
                        has_data = True
                        sampled_non_empty_cells += 1
                        break

                if has_data:
                    break

            if has_data:
                non_empty_sheets += 1

        if sheet_count == 0 or non_empty_sheets == 0:
            return spreadsheet_failure(
                "EMPTY_SPREADSHEET",
                "Spreadsheet does not contain readable data.",
                file_name
            )

        return {
            "status": SUCCESS,
            "document_type": "SPREADSHEET",
            "validation_strategy": "SPREADSHEET_STRUCTURE_VALIDATION",
            "document_quality_score": 100.0,
            "is_readable": True,
            "reupload_required": False,
            "message": "Spreadsheet is readable and acceptable for submission.",
            "total_pages": 0,
            "metadata": {
                "file_name": file_name,
                "file_type": "XLSX",
                "sheet_count": sheet_count,
                "non_empty_sheets": non_empty_sheets,
                "sampled_non_empty_cells": sampled_non_empty_cells
            }
        }

    except Exception as error:
        return spreadsheet_failure(
            "UNREADABLE_SPREADSHEET",
            f"Spreadsheet could not be opened or read: {str(error)}",
            file_name
        )

def validate_word_document(file_bytes: bytes, file_name: str = "") -> Dict[str, Any]:
    file_name_lower = (file_name or "").lower()

    try:
        if file_name_lower.endswith(".doc"):
            return word_failure(
                "UNSUPPORTED_LEGACY_WORD",
                "Legacy .doc files are not supported in this validator. Please upload .docx.",
                file_name
            )

        document = Document(io.BytesIO(file_bytes))

        paragraph_texts = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text and paragraph.text.strip()
        ]

        table_texts = []

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        table_texts.append(cell_text)

        all_texts = paragraph_texts + table_texts
        combined_text = normalize_text(" ".join(all_texts))

        char_count = len(combined_text)
        word_count = len(combined_text.split())

        if char_count < settings.word_min_chars or word_count < settings.word_min_words:
            return word_failure(
                "EMPTY_OR_LOW_CONTENT_WORD_DOCUMENT",
                "Word document does not contain enough readable text.",
                file_name
            )

        return {
            "status": SUCCESS,
            "document_type": "WORD_DOCUMENT",
            "validation_strategy": "WORD_DOCUMENT_STRUCTURE_VALIDATION",
            "document_quality_score": 100.0,
            "is_readable": True,
            "reupload_required": False,
            "message": "Word document is readable and acceptable for submission.",
            "total_pages": 0,
            "metadata": {
                "file_name": file_name,
                "file_type": "DOCX",
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "non_empty_text_blocks": len(all_texts),
                "char_count": char_count,
                "word_count": word_count,
                "sample_text": combined_text[:300]
            }
        }

    except Exception as error:
        return word_failure(
            "UNREADABLE_WORD_DOCUMENT",
            f"Word document could not be opened or read: {str(error)}",
            file_name
        )

def validate_document_quality(
    file_bytes: bytes,
    ocr_service,
    file_name: str = "",
    content_type: str = ""
) -> Dict[str, Any]:

    if not file_bytes:
        raise ValueError("Empty file uploaded.")

    size_mb = len(file_bytes) / (1024 * 1024)

    if size_mb > settings.max_file_size_mb:
        raise ValueError(
            f"File too large. Maximum allowed size is {settings.max_file_size_mb} MB."
        )

    if is_spreadsheet(file_name):
        return make_json_safe(
            validate_spreadsheet(
                file_bytes=file_bytes,
                file_name=file_name
            ))

    if is_word_document(file_name):
        return make_json_safe(
            validate_word_document(
                file_bytes=file_bytes,
                file_name=file_name
            )
        )
    page_results = []

    if is_pdf(file_bytes):
        classification = classify_pdf(file_bytes)
        total_pages = classification["total_pages"]

        scanned_candidates = []
        failed_pages_count = 0

        for page_info in classification["pages"]:
            page_number = page_info["page_number"]
            page_type = page_info["page_type"]
            strategy = get_validation_strategy(page_type)

            if page_type == "DIGITAL_TEXT_PAGE":
                decision = decide_text_layer_page(page_info)

                page_results.append({
                    "page_number": page_number,
                    "page_type": page_type,
                    "validation_strategy": strategy,
                    "status": decision["status"],
                    "quality_score": decision["quality_score"],
                    "issues": decision["issues"],
                    "warnings": decision["warnings"],
                    "metrics": decision["metrics"],
                })

            elif page_type == "OCR_LAYERED_SCAN_PAGE":
                rendered = render_pdf_page(
                    file_bytes=file_bytes,
                    page_index=page_number - 1,
                    dpi=settings.render_dpi
                )

                image_rgb = png_bytes_to_rgb_image(rendered)
                image_quality = analyze_image_quality(image_rgb)

                decision = decide_searchable_scanned_page(
                    page_info=page_info,
                    image_quality_result=image_quality
                )

                page_results.append({
                    "page_number": page_number,
                    "page_type": page_type,
                    "validation_strategy": strategy,
                    "status": decision["status"],
                    "quality_score": decision["quality_score"],
                    "issues": decision["issues"],
                    "warnings": decision["warnings"],
                    "metrics": decision["metrics"],
                })

            elif page_type in ["SCANNED_IMAGE_PAGE", "UNKNOWN_PAGE"]:
                rendered = render_pdf_page(
                    file_bytes=file_bytes,
                    page_index=page_number - 1,
                    dpi=settings.render_dpi
                )

                image_rgb = png_bytes_to_rgb_image(rendered)
                image_quality = analyze_image_quality(image_rgb)

                issues = image_quality["issues"]

                if has_blocking_critical_issue(issues):
                    decision = decide_image_quality_only_page(
                        page_info=page_info,
                        image_quality_result=image_quality
                    )

                    page_results.append({
                        "page_number": page_number,
                        "page_type": page_type,
                        "validation_strategy": strategy,
                        "status": decision["status"],
                        "quality_score": decision["quality_score"],
                        "issues": decision["issues"],
                        "warnings": decision["warnings"],
                        "metrics": decision["metrics"],
                    })

                    if decision["status"] == FAILURE:
                        failed_pages_count += 1

                    if failed_pages_count >= settings.max_failed_pages_before_reject:
                        break

                else:
                    scanned_candidates.append({
                        "page_number": page_number,
                        "page_type": page_type,
                        "validation_strategy": strategy,
                        "page_info": page_info,
                        "image_quality_result": image_quality,
                        "image_rgb": image_rgb,
                        "ocr_recovery_required": requires_ocr_recovery(issues),
                    })
            else:
                decision = decide_blank_or_low_content_page(
                    page_info=page_info,
                    total_pages=total_pages
                )

                page_results.append({
                    "page_number": page_number,
                    "page_type": page_type,
                    "validation_strategy": strategy,
                    "status": decision["status"],
                    "quality_score": decision["quality_score"],
                    "issues": decision["issues"],
                    "warnings": decision["warnings"],
                    "metrics": decision["metrics"],
                })

        selected_pages = select_ocr_candidate_pages(
            scanned_candidates=scanned_candidates,
            max_ocr_pages=settings.max_ocr_pages_per_document
        )

        selected_set = set(selected_pages)

        for candidate in scanned_candidates:
            page_number = candidate["page_number"]
            page_type = candidate["page_type"]
            strategy = candidate["validation_strategy"]
            page_info = candidate["page_info"]
            image_quality = candidate["image_quality_result"]

            if page_number in selected_set:
                image_rgb = candidate["image_rgb"]
                ocr_metrics = ocr_service.run_ocr(image_rgb)

                if candidate.get("ocr_recovery_required"):
                    decision = decide_ocr_recovery_page(
                        image_quality_result=image_quality,
                        ocr_metrics=ocr_metrics
                    )
                else:
                    decision = decide_ocr_page(
                        image_quality_result=image_quality,
                        ocr_metrics=ocr_metrics,
                        page_type=page_type,
                        total_pages=total_pages
                    )

                metrics = {
                    "text_stats": page_info["text_stats"],
                    "image_stats": page_info["image_stats"],
                    "image_quality": decision["metrics"]["image_quality"],
                    "ocr": decision["metrics"]["ocr"],
                }

                metrics = {
                    "text_stats": page_info["text_stats"],
                    "image_stats": page_info["image_stats"],
                    "image_quality": decision["metrics"]["image_quality"],
                    "ocr": decision["metrics"]["ocr"],
                }

            else:
                decision = decide_image_quality_only_page(
                    page_info=page_info,
                    image_quality_result=image_quality
                )

                metrics = decision["metrics"]

            page_results.append({
                "page_number": page_number,
                "page_type": page_type,
                "validation_strategy": strategy,
                "status": decision["status"],
                "quality_score": decision["quality_score"],
                "issues": decision["issues"],
                "warnings": decision["warnings"],
                "metrics": metrics,
            })

            if decision["status"] == FAILURE:
                failed_pages_count += 1

            if failed_pages_count >= settings.max_failed_pages_before_reject:
                break

        document_decision = decide_document(
            page_results=page_results,
            total_pages=total_pages
        )

        return make_json_safe(
            build_ui_response(
                document_decision=document_decision,
                document_type=classification["document_type"],
                validation_strategy="PDF_PAGE_LEVEL_HYBRID_VALIDATION",
                total_pages=total_pages,
                page_results=page_results,
                page_type_counts=classification["page_type_counts"],
            )
        )

    try:
        image_rgb = bytes_to_rgb_image(file_bytes)

    except Exception:
        raise ValueError("Unsupported or corrupted image/document format.")

    image_quality = analyze_image_quality(image_rgb)

    issues = image_quality["issues"]

    if has_blocking_critical_issue(issues):
        decision = {
            "status": FAILURE,
            "quality_score": image_quality["quality_score"],
            "issues": [
                issue for issue in image_quality["issues"]
                if issue.get("severity") == "CRITICAL"
            ],
            "warnings": [
                issue for issue in image_quality["issues"]
                if issue.get("severity") != "CRITICAL"
            ],
            "metrics": {
                "image_quality": image_quality["metrics"],
                "ocr": {
                    "skipped": True,
                    "reason": "OCR skipped because image quality has blocking critical issues.",
                },
            },
        }

    elif requires_ocr_recovery(issues):
        ocr_metrics = ocr_service.run_ocr(image_rgb)
        decision = decide_ocr_recovery_page(
            image_quality_result=image_quality,
            ocr_metrics=ocr_metrics
        )

    else:
        ocr_metrics = ocr_service.run_ocr(image_rgb)
        decision = decide_ocr_page(
            image_quality_result=image_quality,
            ocr_metrics=ocr_metrics,
            page_type="IMAGE_DOCUMENT",
            total_pages=1
        )

    page_results.append({
        "page_number": 1, 
        "page_type": "IMAGE_DOCUMENT",
        "validation_strategy": get_validation_strategy("IMAGE_DOCUMENT"),
        "status": decision["status"],
        "quality_score": decision["quality_score"],
        "issues": decision["issues"],
        "warnings": decision["warnings"],
        "metrics": decision["metrics"],
    })

    document_decision = decide_document(
        page_results=page_results,
        total_pages=1
    )

    return make_json_safe(
        build_ui_response(
            document_decision=document_decision,
            document_type="IMAGE_DOCUMENT",
            validation_strategy="OCR_IMAGE_QUALITY_VALIDATION",
            total_pages=1,
            page_results=page_results,
            page_type_counts={"IMAGE_DOCUMENT": 1},
        )
    )